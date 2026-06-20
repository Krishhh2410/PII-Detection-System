"""
Script to generate sample PDF and DOCX files with PII for testing.
"""
import os
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_sample_pdf(output_path: str):
    """Generate a sample PDF document with PII."""
    doc = SimpleDocTemplate(output_path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=30,
        alignment=1  # Center
    )
    story.append(Paragraph("Employee Information Report", title_style))
    story.append(Spacer(1, 20))
    
    # Introduction
    story.append(Paragraph("Confidential - Internal Use Only", styles['Heading3']))
    story.append(Spacer(1, 10))
    
    intro_text = """
    This document contains sensitive employee information for TechCorp India Pvt Ltd.<br/>
    Generated on: January 15, 2024<br/>
    Document ID: DOC-2024-001
    """
    story.append(Paragraph(intro_text, styles['Normal']))
    story.append(Spacer(1, 20))
    
    # Employee data table
    story.append(Paragraph("Employee Details", styles['Heading2']))
    story.append(Spacer(1, 10))
    
    data = [
        ['Name', 'Email', 'Phone', 'Aadhaar', 'PAN', 'Department'],
        ['Rajesh Kumar', 'rajesh.kumar@techcorp.in', '+91-9876543210', '1234 5678 9012', 'ABCDE1234F', 'Engineering'],
        ['Priya Sharma', 'priya.sharma@techcorp.in', '09876543210', '9876 5432 1098', 'FGHIJ5678K', 'Marketing'],
        ['Amit Patel', 'amit.patel@gmail.com', '+91 87654 32109', '4567 8901 2345', 'KLMNO9012P', 'Sales'],
        ['Sneha Reddy', 'sneha.reddy@yahoo.co.in', '918765432109', '7890 1234 5678', 'PQRST3456U', 'HR'],
    ]
    
    table = Table(data)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    
    story.append(table)
    story.append(Spacer(1, 20))
    
    # Individual employee details
    story.append(Paragraph("Detailed Employee Information", styles['Heading2']))
    story.append(Spacer(1, 10))
    
    employees = [
        {
            'name': 'Vikram Singh',
            'email': 'vikram.singh@outlook.com',
            'phone': '080-12345678',
            'aadhaar': '2345 6789 0123',
            'pan': 'UVWXY7890Z',
            'address': '654 Anna Salai, Chennai, Tamil Nadu 600002',
            'dob': '12 September 1983',
            'ip': '10.10.10.10'
        },
        {
            'name': 'Ananya Gupta',
            'email': 'ananya.gupta@customer.in',
            'phone': '+91-9988776655',
            'aadhaar': '3456 7890 1234',
            'pan': 'ABCPQ1234D',
            'address': '45 Salt Lake, Kolkata, West Bengal 700091',
            'dob': '20 May 1995',
            'ip': '203.0.113.50'
        }
    ]
    
    for emp in employees:
        emp_text = f"""
        <b>Name:</b> {emp['name']}<br/>
        <b>Email:</b> {emp['email']}<br/>
        <b>Phone:</b> {emp['phone']}<br/>
        <b>Aadhaar:</b> {emp['aadhaar']}<br/>
        <b>PAN:</b> {emp['pan']}<br/>
        <b>Address:</b> {emp['address']}<br/>
        <b>Date of Birth:</b> {emp['dob']}<br/>
        <b>Last Login IP:</b> {emp['ip']}<br/>
        """
        story.append(Paragraph(emp_text, styles['Normal']))
        story.append(Spacer(1, 15))
    
    # Contact information
    story.append(Paragraph("Emergency Contacts", styles['Heading2']))
    story.append(Spacer(1, 10))
    
    contact_text = """
    <b>Primary Contact:</b> Suresh Kumar - +91-9876500001<br/>
    <b>Secondary Contact:</b> Meera Devi - 919876500002<br/>
    <b>Email:</b> contact@example.in, info@company.co.in<br/>
    """
    story.append(Paragraph(contact_text, styles['Normal']))
    
    # Build PDF
    doc.build(story)
    print(f"PDF generated: {output_path}")


def generate_sample_docx(output_path: str):
    """Generate a sample DOCX document with PII."""
    doc = Document()
    
    # Title
    title = doc.add_heading('Customer Data Export', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Header info
    doc.add_paragraph('Confidential Document')
    doc.add_paragraph('Company: TechCorp India Pvt Ltd')
    doc.add_paragraph('Generated: January 15, 2024')
    doc.add_paragraph()
    
    # Section 1: Customer List
    doc.add_heading('Customer Contact Information', level=1)
    
    # Create table
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    hdr_cells = table.rows[0].cells
    headers = ['Full Name', 'Email Address', 'Phone Number', 'Aadhaar', 'PAN', 'Credit Card']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    
    # Add customer data
    customers = [
        ('Rahul Verma', 'rahul.verma@email.com', '919988776655', '5678 9012 3456', 'XYZAB5678C', '4111-1111-1111-1111'),
        ('Ananya Gupta', 'ananya.gupta@customer.in', '+91-9988776655', '3456 7890 1234', 'ABCPQ1234D', '4532-1234-5678-9012'),
        ('Rajesh Kumar', 'rajesh.kumar@techcorp.in', '+91-9876543210', '1234 5678 9012', 'ABCDE1234F', '3712-345678-90123'),
    ]
    
    for customer in customers:
        row_cells = table.add_row().cells
        for i, value in enumerate(customer):
            row_cells[i].text = value
    
    doc.add_paragraph()
    
    # Section 2: Detailed Profiles
    doc.add_heading('Detailed Customer Profiles', level=1)
    
    profiles = [
        {
            'name': 'Rahul Verma',
            'email': 'rahul.verma@email.com',
            'phone': '919988776655',
            'alt_phone': '+91-11-12345678',
            'aadhaar': '5678 9012 3456',
            'pan': 'XYZAB5678C',
            'address': '12 Jubilee Hills, Hyderabad 500033',
            'office': 'Madhapur, Hyderabad 500081',
            'dob': '03 December 1987',
            'ip': '198.51.100.25'
        },
        {
            'name': 'Priya Sharma',
            'email': 'priya.sharma@techcorp.in',
            'phone': '09876543210',
            'alt_phone': '022-87654321',
            'aadhaar': '9876 5432 1098',
            'pan': 'FGHIJ5678K',
            'address': '456 Park Street, Mumbai, Maharashtra 400001',
            'office': 'Bandra Kurla Complex, Mumbai 400051',
            'dob': '22 July 1990',
            'ip': '10.0.0.50'
        }
    ]
    
    for profile in profiles:
        doc.add_heading(profile['name'], level=2)
        
        p = doc.add_paragraph()
        p.add_run('Email: ').bold = True
        p.add_run(profile['email'])
        
        p = doc.add_paragraph()
        p.add_run('Phone: ').bold = True
        p.add_run(profile['phone'])
        
        p = doc.add_paragraph()
        p.add_run('Alternate Phone: ').bold = True
        p.add_run(profile['alt_phone'])
        
        p = doc.add_paragraph()
        p.add_run('Aadhaar: ').bold = True
        p.add_run(profile['aadhaar'])
        
        p = doc.add_paragraph()
        p.add_run('PAN: ').bold = True
        p.add_run(profile['pan'])
        
        p = doc.add_paragraph()
        p.add_run('Residential Address: ').bold = True
        p.add_run(profile['address'])
        
        p = doc.add_paragraph()
        p.add_run('Office Address: ').bold = True
        p.add_run(profile['office'])
        
        p = doc.add_paragraph()
        p.add_run('Date of Birth: ').bold = True
        p.add_run(profile['dob'])
        
        p = doc.add_paragraph()
        p.add_run('Registered IP: ').bold = True
        p.add_run(profile['ip'])
        
        doc.add_paragraph()
    
    # Section 3: System Access Logs
    doc.add_heading('Recent System Access', level=1)
    
    access_logs = [
        ('admin@techcorp.in', '192.168.1.1', 'Accessed employee records', '2024-01-15 10:30:00'),
        ('manager@techcorp.in', '10.0.0.100', 'Downloaded customer report', '2024-01-15 14:45:00'),
        ('rajesh.kumar@techcorp.in', '172.16.0.50', 'Updated profile', '2024-01-15 16:20:00'),
    ]
    
    log_table = doc.add_table(rows=1, cols=4)
    log_table.style = 'Light List Accent 1'
    
    hdr_cells = log_table.rows[0].cells
    log_headers = ['User Email', 'IP Address', 'Action', 'Timestamp']
    for i, header in enumerate(log_headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    
    for log in access_logs:
        row_cells = log_table.add_row().cells
        for i, value in enumerate(log):
            row_cells[i].text = value
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph('This document contains confidential information. Do not distribute.')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.italic = True
    footer.runs[0].font.size = Pt(9)
    
    # Save document
    doc.save(output_path)
    print(f"DOCX generated: {output_path}")


if __name__ == "__main__":
    # Get the directory of this script
    script_dir = os.path.dirname(os.path.abspath(__file__))
    
    # Generate sample files
    pdf_path = os.path.join(script_dir, "sample_document.pdf")
    docx_path = os.path.join(script_dir, "sample_document.docx")
    
    generate_sample_pdf(pdf_path)
    generate_sample_docx(docx_path)
    
    print("\nSample files generated successfully!")
    print(f"- SQL file: {os.path.join(script_dir, 'sample_data.sql')}")
    print(f"- PDF file: {pdf_path}")
    print(f"- DOCX file: {docx_path}")
