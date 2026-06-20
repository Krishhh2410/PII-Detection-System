import io
from typing import List
from docx import Document
from docx.shared import RGBColor

from app.services.sanitizers.base import BaseSanitizer
from app.services.pii_detector import PIIDetection


class DOCXSanitizer(BaseSanitizer):
    """Sanitizer for DOCX files."""
    
    def sanitize(self, content: bytes, detections: List[PIIDetection]) -> bytes:
        """
        Sanitize DOCX content by replacing detected PII in text.
        """
        # Load document from bytes
        doc = Document(io.BytesIO(content))
        
        # Sort detections by position (reverse for replacement)
        sorted_detections = sorted(detections, key=lambda d: d.start, reverse=True)
        
        # Process all paragraphs
        for para in doc.paragraphs:
            self._sanitize_paragraph(para, sorted_detections)
        
        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._sanitize_paragraph(para, sorted_detections)
        
        # Process headers and footers
        for section in doc.sections:
            for para in section.header.paragraphs:
                self._sanitize_paragraph(para, sorted_detections)
            for para in section.footer.paragraphs:
                self._sanitize_paragraph(para, sorted_detections)
        
        # Save to bytes
        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()
    
    def _sanitize_paragraph(self, para, detections: List[PIIDetection]):
        """Sanitize text in a paragraph."""
        if not para.text:
            return
        
        full_text = para.text
        modified = False
        
        # Check if any detection text exists in this paragraph
        for detection in detections:
            if detection.text in full_text:
                sanitized = self.apply_anonymization(detection.text, detection)
                full_text = full_text.replace(detection.text, sanitized)
                modified = True
        
        if modified:
            # Clear the paragraph and add new text
            # Note: This is a simplified approach - preserving formatting is complex
            para.clear()
            run = para.add_run(full_text)
    
    def extract_text(self, content: bytes) -> str:
        """Extract text from DOCX for PII detection."""
        doc = Document(io.BytesIO(content))
        text_parts = []
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            text_parts.append(para.text)
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_parts.append(cell.text)
        
        return "\n".join(text_parts)
